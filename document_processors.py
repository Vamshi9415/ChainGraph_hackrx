import os
import uuid
import logging
import asyncio
import requests
from io import BytesIO
from typing import List, Dict, Any, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from urllib.parse import urlparse
from bs4 import BeautifulSoup
import pytesseract

import fitz  # PyMuPDF for PDFs
import docx
import pandas as pd
from PIL import Image
from langchain_community.document_loaders import WebBaseLoader

# Import your custom modules
from config import OCR_AVAILABLE, TEMP_FILES_PATH
from data_models import ExtractedURL, ExtractedTable, ExtractedImage, ProcessedDocument
from language_utils import detect_language_robust, get_language_name

from url_extractor import URLExtractor
from pdf_processor import EnhancedPDFProcessor

from xlsx_table_extractor import EnhancedXLSXTableExtractor

from pptx_processor import AdvancedPPTXProcessor

# --- Enhanced Image OCR Processor ---
class ImageOCRProcessor:
    """Specialized processor for PNG/JPEG images using Pytesseract"""

    @staticmethod
    async def process_image_file(file_content: bytes, file_path: str, request_id: str) -> List[ExtractedImage]:
        """Process images with Pytesseract OCR"""
        images = []
        if not OCR_AVAILABLE:
            logging.error("Pytesseract not available but image processing requested")
            return images

        temp_img_path = os.path.join(TEMP_FILES_PATH, f"{request_id}_{uuid.uuid4().hex}.png")

        try:
            pil_image = Image.open(BytesIO(file_content))

            if pil_image.mode != 'RGB':
                pil_image = pil_image.convert('RGB')

            pil_image.save(temp_img_path, 'PNG')
            width, height = pil_image.size

            try:
                ocr_data = pytesseract.image_to_data(pil_image, output_type=pytesseract.Output.DATAFRAME, lang='eng')
                ocr_data = ocr_data[ocr_data.conf > 0]

                if not ocr_data.empty:
                    final_ocr_text = " ".join(ocr_data['text'].dropna().astype(str))
                    average_confidence = ocr_data['conf'].mean() / 100.0

                    images.append(ExtractedImage(
                        image_path=temp_img_path,
                        ocr_text=final_ocr_text.strip(),
                        metadata={
                            'source': file_path,
                            'extraction_method': 'pytesseract',
                            'image_dimensions': f"{width}x{height}",
                            'processing_timestamp': datetime.now().isoformat(),
                            'mean_confidence_score': f"{average_confidence:.2f}"
                        },
                        confidence=average_confidence
                    ))
                    logging.info(f"Successfully extracted text from image with Pytesseract: {len(final_ocr_text)} characters")
                else:
                    logging.warning("No text with sufficient confidence extracted from image using Pytesseract")

            except pytesseract.TesseractNotFoundError:
                logging.error("Tesseract is not installed or not in your PATH. OCR will not work.")
            except Exception as e:
                logging.error(f"Pytesseract OCR processing failed: {e}")

        except Exception as e:
            logging.error(f"Image processing failed for {file_path}: {e}")
        finally:
            if os.path.exists(temp_img_path):
                try:
                    os.remove(temp_img_path)
                except Exception as e:
                    logging.warning(f"Failed to cleanup temp image file: {e}")

        return images

# --- Enhanced Document Processor ---
class TargetedDocumentProcessor:
    """Enhanced document processor with comprehensive type handling"""

    @staticmethod
    async def process_document(doc_content: bytes, doc_type: str, doc_url: str, request_id: str) -> ProcessedDocument:
        """Process documents with enhanced type-specific handlers"""
        text = ""
        tables = []
        images = []
        extracted_urls = []
        dataframes = {}

        loop = asyncio.get_running_loop()

        try:
            with ThreadPoolExecutor(max_workers=2) as executor:
                if doc_type == "pdf":
                    text = await loop.run_in_executor(
                        executor,
                        EnhancedPDFProcessor.extract_pdf_content,
                        doc_content
                    )

                elif doc_type == "docx":
                    text = await loop.run_in_executor(
                        executor,
                        TargetedDocumentProcessor._extract_docx_text,
                        doc_content
                    )

                elif doc_type == "pptx":
                    # Use the enhanced PPTX processor
                    text, extracted_images = await AdvancedPPTXProcessor.process_pptx(
                        doc_content,
                        request_id
                    )
                    images.extend(extracted_images)

                elif doc_type == "xlsx":
                    tables = await loop.run_in_executor(
                        executor,
                        EnhancedXLSXTableExtractor.extract_tables_from_xlsx,
                        doc_content
                    )

                    # Extract dataframes from tables
                    for table_idx, table in enumerate(tables):
                        if table.dataframe is not None:
                            table_id = f"{table.metadata.get('sheet_name', 'sheet')}_{table_idx+1}"
                            dataframes[table_id] = table.dataframe

                    text = "Spreadsheet document with structured data tables."

                elif doc_type in ["png", "jpeg"]:
                    images = await ImageOCRProcessor.process_image_file(
                        doc_content,
                        doc_url,
                        request_id
                    )
                    text = "Image document processed with OCR."

                elif doc_type == "html":
                    try:
                        docs = await loop.run_in_executor(
                            executor,
                            lambda: WebBaseLoader(web_paths=[doc_url]).load()
                        )
                        if docs:
                            text = docs[0].page_content
                        else:
                            response = await loop.run_in_executor(
                                executor,
                                lambda: requests.get(doc_url, timeout=30)
                            )
                            soup = BeautifulSoup(response.text, 'html.parser')
                            for script in soup(["script", "style", "nav", "footer", "header"]):
                                script.decompose()
                            text = soup.get_text(separator=' ', strip=True)
                    except Exception as e:
                        logging.error(f"Failed to process HTML: {e}")
                        text = f"Error processing HTML: {str(e)}"

                elif doc_type == "txt":
                    try:
                        text = doc_content.decode('utf-8', errors='replace')
                    except UnicodeDecodeError:
                        for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                            try:
                                text = doc_content.decode(encoding, errors='replace')
                                break
                            except:
                                pass

                else:
                    text = f"Unsupported document type: {doc_type}"
                    logging.warning(f"Unsupported document type: {doc_type}")

        except Exception as e:
            logging.error(f"Document processing failed for type {doc_type}: {e}")
            text = f"Error processing {doc_type} document: {str(e)}"

        # Extract URLs
        extracted_urls = URLExtractor.extract_urls(text)

        # Combine content
        enhanced_content = text

        # Add table content
        if tables:
            table_content = "\n\n".join([
                f"--- TABLE FROM {table.location} ---\n{table.content}"
                for table in tables
            ])
            enhanced_content += f"\n\n=== EXTRACTED TABLES ===\n{table_content}"

        # Add image content
        if images:
            image_content = "\n\n".join([
                f"--- IMAGE: {img.metadata.get('source', 'unknown')} ---\n{img.ocr_text}"
                for img in images
            ])
            enhanced_content += f"\n\n=== TEXT FROM IMAGES (OCR) ===\n{image_content}"

        # Add URL summary
        if extracted_urls:
            url_content = "\n".join([
                f"URL: {url.url} (Type: {url.url_type})"
                for url in extracted_urls
            ])
            enhanced_content += f"\n\n=== EXTRACTED URLS ===\n{url_content}"

        # Detect language
        detected_language = detect_language_robust(enhanced_content)
        language_name = get_language_name(detected_language)
        logging.info(f"Detected document language: {language_name} ({detected_language})")

        return ProcessedDocument(
            content=enhanced_content,
            metadata={
                "doc_type": doc_type,
                "doc_url": doc_url,
                "char_count": len(enhanced_content),
                "table_count": len(tables),
                "image_count": len(images),
                "url_count": len(extracted_urls),
                "language": language_name,
                "language_code": detected_language,
                "processing_timestamp": datetime.now().isoformat()
            },
            tables=tables,
            images=images,
            extracted_urls=extracted_urls,
            detected_language=detected_language,
            dataframes=dataframes
        )

    @staticmethod
    def _extract_docx_text(doc_content: bytes) -> str:
        """Enhanced DOCX text extraction"""
        text_parts = []

        try:
            doc = docx.Document(BytesIO(doc_content))

            core_properties = doc.core_properties
            if core_properties:
                prop_parts = []
                if core_properties.title:
                    prop_parts.append(f"Title: {core_properties.title}")
                if core_properties.author:
                    prop_parts.append(f"Author: {core_properties.author}")
                if core_properties.subject:
                    prop_parts.append(f"Subject: {core_properties.subject}")

                if prop_parts:
                    text_parts.append("=== DOCUMENT PROPERTIES ===")
                    text_parts.extend(prop_parts)
                    text_parts.append("")

            current_heading = None
            paragraph_buffer = []

            for para in doc.paragraphs:
                if not para.text.strip():
                    continue

                if para.style.name.startswith('Heading'):
                    if paragraph_buffer:
                        if current_heading:
                            text_parts.append(f"=== {current_heading} ===")
                        text_parts.append("\n".join(paragraph_buffer))
                        text_parts.append("")
                        paragraph_buffer = []

                    current_heading = para.text.strip()
                else:
                    paragraph_buffer.append(para.text.strip())

            if paragraph_buffer:
                if current_heading:
                    text_parts.append(f"=== {current_heading} ===")
                text_parts.append("\n".join(paragraph_buffer))

            for table_idx, table in enumerate(doc.tables, 1):
                table_text = []
                table_text.append(f"=== TABLE {table_idx} ===")

                for row_idx, row in enumerate(table.rows):
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        prefix = "HEADER: " if row_idx == 0 else f"ROW {row_idx}: "
                        table_text.append(f"{prefix}{row_text}")

                if len(table_text) > 1:
                    text_parts.append("\n".join(table_text))
                    text_parts.append("")

            logging.info(f"Extracted text from DOCX: {len(text_parts)} elements")

        except Exception as e:
            logging.error(f"DOCX extraction failed: {e}")
            text_parts.append(f"Error extracting DOCX content: {str(e)}")

        return "\n".join(text_parts)


